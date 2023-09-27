'''
FreeBingo+.
Copyright (C) 2021 by Fedor Egorov
This file is part of FreeBingo+.

FreeBingo+ is free software: you can redistribute it and/or modify
it under the terms of the GNU General Public License as published by
the Free Software Foundation, either version 3 of the License, or
(at your option) any later version.

FreeBingo+ is distributed in the hope that it will be useful,
but WITHOUT ANY WARRANTY; without even the implied warranty of
MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
GNU General Public License for more details.

You should have received a copy of the GNU General Public License
along with FreeBingo+. If not, see <https://www.gnu.org/licenses/>.
'''

from docx import Document;
from docx.shared import Pt;
from datetime import datetime;
from tkinter import *;
from tkinter import ttk;
from tkinter.filedialog import *;
from PIL import ImageTk, Image;
from webbrowser import open_new;
import ctypes;
import sys;
from subprocess import call;
from random import randint;

def open_site(url):
    open_new(url);

def on_close():
    sys.exit();

def gen():
    docx = Document('resources/template.docx');

    for table in docx.tables: # Read all tables
        n = 1; # "ID" of cell
        for row in table.rows: # Read all rows
            for cell in row.cells: # Read all cells
                if cell.text == '' and n != 23: # Check for letters and "ID"
                    cell.text = f'\n{randint(0, 100)}'; # Paste random number
                    paragraphs = cell.paragraphs; # Get paragraphs
                    for paragraph in paragraphs: # Read all paragraphs
                        for run in paragraph.runs: # Read all runs
                            font = run.font; # Get font
                            font.name = 'Arial'; # Change style of font
                            run.font.bold = True; # Make font bold
                        paragraph.alignment = 1; # Center parapraph
                n += 1; # Next "ID"

    for paragraph in docx.paragraphs:
        if paragraph.text == 'Generated: ':
            paragraph.text += datetime.now(
                ).strftime("%D %H:%M (month/day/year)");
            for run in paragraph.runs:
                run.font.name = 'Arial';
                run.font.size = Pt(11);
                run.font.bold = True;
                run.font.italic = True;

    p = asksaveasfilename(defaultextension = '.docx',
        filetypes = [("Document file", '*.docx')], initialfile = '*.docx');
    if p != '':
        docx.save(p);

try:
    # Kill task
    call('taskkill /f /im freebingoplus.exe', shell = True);

    # FOR ICON ON PANEL
    myappid = 'fedoregorov.freebingoplus.1.1';
    ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid);

    # INITIALIZE WINDOW
    Tk().withdraw();
    window = Toplevel();

    # CENTER WINDOW
    screen_width = window.winfo_screenwidth();
    screen_height = window.winfo_screenheight();

    x_cordinate = int((screen_width / 2) - (500 / 2));
    y_cordinate = int((screen_height / 2) - (200 / 2));
    window.geometry("{}x{}+{}+{}".format(500, 200, x_cordinate, y_cordinate));
    window.resizable(width = False, height = False);

    # ICON & TITLE
    window.iconbitmap('resources/icon.ico');
    window.title('FreeBingo+');

    # INITIALIZE NOTEBOOK (TABS) AND FRAMES
    nb = ttk.Notebook(window);
    f1 = Frame(window); # Generate bingo frame
    f2 = Frame(window); # About program frame

    nb.add(f1, text = 'Generate bingo'); # Link frames and tabs
    nb.add(f2, text = 'About');
    nb.pack(fill = 'both', expand = 'yes');

    # GENERATE BINGO FRAME
    Label(f1, text = 'Welcome!', font = ('Segoe UI',
        16, 'bold')).place(x = 190, y = 30); # Welcome label

    Label(f1, text = 'FreeBingo+', font = ('Segoe UI', 32, 'bold'),
        fg='#ccc').place(x = 270, y = 120); # 'FreeBingo+' label

    genb = ttk.Button(f1, text = 'Generate!', command = gen);
    genb.place(x = 203, y = 75);

    # ABOUT PROGRAM FRAME
    Label(f2, text = 'FreeBingo+', font = ('Segoe UI', 32, 'bold'),
        fg='#ccc').place(x = 270, y = 120); # 'FreeBingo+' label

    img = Image.open('resources/icon.png'); # Open image and resize
    img = img.resize((73, 73), Image.ANTIALIAS);
    img = ImageTk.PhotoImage(img); # Initialize PhotoImage from PIL and place
    Label(f2, text = 'Img', image = img).place(x = 30, y = 30);

    Label(f2, text = 'FreeBingo+', font = ('Segoe UI',
        10, 'bold')).place(x = 110, y = 30);
    
    Label(f2, text = 'Version 1.1 (2021)').place(x = 110, y = 50);
    Label(f2, text = 'Author: Fedor Egorov (FatlessComb1168)').place(
        x = 110, y = 67);

    gh = Label(f2, text = '(GitHub)', fg = "blue",
        cursor = "hand2", underline = True);
    gh.place(x = 330, y = 67);
    gh.bind("<Button-1>", lambda e:
        open_site("https://github.com/FatlessComb1168"));

    Label(f2, text = 'Distributed under GNU GPL 3.0.').place(x = 110, y = 87);
    lsc = Label(f2, text = 'More details...', fg = "blue",
        cursor = "hand2", underline = True);
    lsc.place(x = 280, y = 87);
    lsc.bind("<Button-1>", lambda e:
        open_site("https://www.gnu.org/licenses/gpl-3.0.en.html"));

    # CHECK FOR CLOSING
    window.protocol('WM_DELETE_WINDOW', on_close);

    # MAINLOOP
    window.mainloop();

except Exception as e:
    print('Error:', e);
    input();
